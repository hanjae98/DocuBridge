import os
import threading
import datetime
import re
import time
import random
import json
import webbrowser
import platform
import subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext
import requests
import atexit

from packaging import version
from docx import Document
from docx.shared import RGBColor, Pt
import translators as ts
from concurrent.futures import ThreadPoolExecutor, wait, FIRST_COMPLETED

# ===== [Settings] =====
APP_NAME = "DocuBridge"
APP_SUBTITLE = "Hybrid (Cloud + On-Device) Translator"
CURRENT_VERSION = "v1.1.0"
REPO_OWNER = "hanjae98" 
REPO_NAME = "DocuBridge"    
VERSION_URL = f"https://raw.githubusercontent.com/{REPO_OWNER}/{REPO_NAME}/main/version.json"
RELEASE_URL = f"https://github.com/{REPO_OWNER}/{REPO_NAME}/releases/latest"

MOCK_TEST = False 
APPEND_COLOR = (128, 128, 128)
# Hybrid 모드에서는 Online이 빠르므로 4로 유지하되, Local 사용 시 내부 Lock으로 제어됨
MAX_WORKERS = 4 

HAN_TO_ENG_MAP = {
    '가': 'A', '나': 'B', '다': 'C', '라': 'D', '마': 'E', '바': 'F', '사': 'G',
    '아': 'H', '자': 'I', '차': 'J', '카': 'K', '타': 'L', '파': 'M', '하': 'N',
    '①': '1', '②': '2', '③': '3', '④': '4', '⑤': '5', '⑥': '6', '⑦': '7',
    '⑧': '8', '⑨': '9', '⑩': '10', '⑪': '11', '⑫': '12', '⑬': '13', '⑭': '14', '⑮': '15',
    'ㄱ': 'a', 'ㄴ': 'b', 'ㄷ': 'c', 'ㄹ': 'd', 'ㅁ': 'e', 'ㅂ': 'f', 'ㅅ': 'g',
    'ㅇ': 'h', 'ㅈ': 'i', 'ㅊ': 'j', 'ㅋ': 'k', 'ㅌ': 'l', 'ㅍ': 'm', 'ㅎ': 'n'
}

# ===== [Helper(Config)] =====
class ConfigManager:
    def __init__(self):
        self.config_file = "config.json"
        self.defaults = {
            "ignored_version": "v0.0.0",
            "theme": "light",
            "debug_mode": False,
            "backend_priority": "online",  # 'online' or 'local'
            "ollama_model": "qwen2.5:1.5b" # Default AI Model
        }
        self.data = self.load()
        atexit.register(self.save)

    def load(self):
        if not os.path.exists(self.config_file):
            return self.defaults.copy()
        try:
            with open(self.config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return self.defaults.copy()

    def save(self):
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.data, f, indent=4, ensure_ascii=False)
        except Exception as e:
            print(f"Config save failed: {e}")

    def get(self, key, default=None):
        return self.data.get(key, self.defaults.get(key, default))

    def set(self, key, value):
        self.data[key] = value
        self.save()

config = ConfigManager()

class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip_window = None
        self.widget.bind("<Enter>", self.show_tip)
        self.widget.bind("<Leave>", self.hide_tip)
    def show_tip(self, event=None):
        if self.tip_window or not self.text: return
        x, y, _, _ = self.widget.bbox("insert")
        x = x + self.widget.winfo_rootx() + 25
        y = y + self.widget.winfo_rooty() + 25
        self.tip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(tw, text=self.text, justify='left',
                        background="#ffffe0", relief='solid', borderwidth=1,
                        font=("Segoe UI", 9))
        label.pack(ipadx=1)
    def hide_tip(self, event=None):
        if self.tip_window:
            self.tip_window.destroy()
            self.tip_window = None

class LifecycleManager:
    def __init__(self):
        self.tracking_table = {} 
        self.lock = threading.Lock()
    def register(self, task_id, original_text):
        with self.lock: self.tracking_table[task_id] = {"status": "READY", "result": None, "orig": original_text}
    def update_status(self, task_id, status, result=None):
        with self.lock:
            self.tracking_table[task_id]["status"] = status
            if result: self.tracking_table[task_id]["result"] = result
    def get_failed_tasks(self):
        failed = []
        with self.lock:
            for tid, info in self.tracking_table.items():
                if info["status"] == "FAILED": failed.append((tid, info["orig"]))
        return failed
    def get_summary(self):
        summary = {"SUCCESS": 0, "SKIPPED": 0, "FAILED": 0, "READY": 0, "IN_PROGRESS": 0}
        with self.lock:
            for info in self.tracking_table.values():
                s = info["status"]
                summary[s] = summary.get(s, 0) + 1
        return summary

lifecycle_manager = LifecycleManager()

class UpdateManager:
    def __init__(self, app_instance):
        self.app = app_instance
        self.latest_version = None
        self.update_available = False
    def check_for_updates(self):
        if MOCK_TEST:
            time.sleep(2)
            self.update_available = True
            self.app.notify_update("v9.9.9", "⭐ Mock update test.")
            return
        try:
            response = requests.get(VERSION_URL, timeout=3)
            if response.status_code == 200:
                data = response.json()
                remote_ver = data.get("version", "v0.0.0")
                if version.parse(remote_ver) > version.parse(CURRENT_VERSION):
                    if config.get("ignored_version") != remote_ver:
                        self.update_available = True
                        self.app.notify_update(remote_ver, data.get("message", "New features available"))
        except: pass

def get_unique_filename(path, suffix):
    folder, name = os.path.split(path)
    base, ext = os.path.splitext(name)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    new_name = f"{base}_{suffix}_{timestamp}{ext}"
    return os.path.join(folder, new_name)

def open_file_safe(path):
    try:
        if platform.system() == 'Windows': os.startfile(path)
        elif platform.system() == 'Darwin': subprocess.call(('open', path))
        else: subprocess.call(('xdg-open', path))
    except Exception as e: print(f"Error opening file: {e}")

def is_korean_present(text): return any('\uac00' <= c <= '\ud7a3' for c in text)

def is_already_translated_strict(text):
    text = text.strip()
    if not text.endswith(')'): return False
    last_open = text.rfind('(')
    if last_open == -1: return False
    content = text[last_open+1:-1].strip()
    if any('\uac00' <= c <= '\ud7a3' for c in content): return False
    if not any('a' <= c.lower() <= 'z' for c in content): return False
    return True

class FileLogger:
    def __init__(self, filename):
        self.filename = filename
        self.logs = []
        self.lock = threading.Lock()
    def add(self, task_id, status, engine, original, translated):
        with self.lock: self.logs.append({'id': task_id, 'status': status, 'engine': engine, 'orig': original.strip(), 'trans': str(translated).strip()})
    def save(self):
        self.logs.sort(key=lambda x: x['id'])
        with open(self.filename, 'w', encoding='utf-8') as f:
            f.write(f"=== {APP_NAME} Log ({datetime.datetime.now()}) ===\n\n")
            for log in self.logs: f.write(f"[{log['id']:03d}] [{log['status']}] [{log['engine']}]\nORIGIN: {log['orig']}\nTRANS : {log['trans']}\n{'-'*60}\n")
        return self.filename

# ===== [Logic - Translation Backends] =====
class TranslationBackend:
    def check_health(self, app):
        raise NotImplementedError
    def translate(self, text, task_index, app, logger, task_id):
        raise NotImplementedError
    def recover_batch(self, text):
        raise NotImplementedError

# 1. Online Backend (from Old Docubridge)
class OnlineBackend(TranslationBackend):
    def __init__(self):
        self.candidate_engines = ['google', 'bing', 'alibaba']
        self.active_engines = []

    def check_health(self, app):
        self.active_engines = []
        test_text = "테스트"
        for engine in self.candidate_engines:
            try:
                # 짧은 타임아웃으로 상태 확인
                res = ts.translate_text(test_text, translator=engine, from_language='ko', to_language='en', timeout=3)
                if res: self.active_engines.append(engine)
            except: pass
        
        # 최소한 구글은 추가 (실패하더라도 시도는 하도록)
        if not self.active_engines: self.active_engines.append('google')

    def translate(self, text, task_index, app, logger, task_id):
        if not self.active_engines: return None
        
        primary_idx = task_index % len(self.active_engines)
        # Round-robin queue
        queue = [self.active_engines[primary_idx]] + [e for e in self.active_engines if e != self.active_engines[primary_idx]]
        
        if text.strip() == "기타":
            if logger: logger.add(task_id, "REPLACE", "System", text, "Etc")
            return "Etc"
            
        for engine in queue:
            try:
                if app.debug_mode: time.sleep(random.uniform(0.1, 0.3))
                res = ts.translate_text(text, translator=engine, from_language='ko', to_language='en', timeout=5)
                if res:
                    if logger: logger.add(task_id, "SUCCESS", f"Online({engine})", text, res)
                    return res
            except: continue
        return None

    def recover_batch(self, text):
        # Recovery Logic: Try all engines concurrently
        with ThreadPoolExecutor(max_workers=len(self.active_engines)) as executor:
            futures = {executor.submit(ts.translate_text, text, translator=eng, from_language='ko', to_language='en'): eng for eng in self.active_engines}
            done, not_done = wait(futures, return_when=FIRST_COMPLETED)
            for future in done:
                try:
                    result = future.result()
                    if result: return result, futures[future]
                except: continue
        return None, None

# 2. Local AI Backend (Ollama)
class OllamaBackend(TranslationBackend):
    def __init__(self):
        self.api_url = "http://localhost:11434/api/generate"
        self.model_name = config.get("ollama_model", "qwen2.5:1.5b")
        # CPU/GPU 리소스 보호를 위해 Lock 사용 (MAX_WORKERS=4여도 Ollama는 1개씩 or 병렬설정따라)
        self.lock = threading.Lock()
        self.is_available = False

    def check_health(self, app):
        # 1. Ollama 실행 여부 확인
        try:
            requests.get("http://localhost:11434/", timeout=2)
        except:
            # Ollama가 꺼져있으면 Local 사용 불가 처리
            self.is_available = False
            return

        # 2. 모델 존재 여부 확인 및 자동 설치
        try:
            res = requests.get("http://localhost:11434/api/tags", timeout=5)
            models = [m['name'] for m in res.json().get('models', [])]
            
            if not any(self.model_name in m for m in models):
                # 모델이 없으면 사용자에게 물어보고 다운로드
                if messagebox.askyesno("AI Model Missing", 
                                       f"로컬 AI 모델 '{self.model_name}'이 없습니다.\n다운로드하시겠습니까? (약 1~2GB)"):
                    self.download_model(app)
                else:
                    self.is_available = False
                    return
            
            self.is_available = True
            
        except Exception as e:
            self.is_available = False
            if app.debug_mode: print(f"Ollama Health Check Error: {e}")

    def download_model(self, app):
        app.update_status_text(f"Downloading {self.model_name}... (This may take a while)")
        # 윈도우 터미널을 열지 않고 백그라운드에서 실행
        try:
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            process = subprocess.Popen(["ollama", "pull", self.model_name], 
                                       stdout=subprocess.PIPE, stderr=subprocess.PIPE, 
                                       startupinfo=startupinfo)
            process.wait() # 다운로드 끝날 때까지 대기
            app.update_status_text("Model download complete.")
        except Exception as e:
            messagebox.showerror("Download Failed", f"모델 다운로드 실패: {e}")

    def translate(self, text, task_index, app, logger, task_id):
        if not self.is_available: return None

        prompt = f"Translate this Korean text to English. Output ONLY the translated text without any explanation.\n\nKorean: {text}\nEnglish:"
        payload = {
            "model": self.model_name,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": 0.0, "num_predict": 128, "num_ctx": 2048}
        }

        with self.lock: # 리소스 보호
            try:
                response = requests.post(self.api_url, json=payload, timeout=60)
                if response.status_code == 200:
                    res_json = response.json()
                    translated = res_json.get("response", "").strip()
                    # 후처리
                    if translated.lower().startswith("english:"): translated = translated[8:].strip()
                    translated = translated.strip('"').strip("'")
                    
                    if translated:
                        if logger: logger.add(task_id, "SUCCESS", "Local_AI", text, translated)
                        return translated
            except Exception as e:
                if logger: logger.add(task_id, "ERROR", "Local_AI", text, str(e))
        return None

    def recover_batch(self, text):
        return self.translate(text, 0, None, None, -1), "Local_AI_Retry"

# 3. Hybrid Manager (The Brain)
class HybridBackendManager:
    def __init__(self):
        self.online = OnlineBackend()
        self.local = OllamaBackend()
        # config.json에서 우선순위 로드 (기본값: online)
        self.priority = config.get("backend_priority", "online") 

    def check_health(self, app):
        app.start_checking_animation()
        
        # 1. Check Primary
        if self.priority == "online":
            app.update_status_text("Checking Online Translators...")
            self.online.check_health(app)
            # 온라인이 불안하면 로컬도 체크해둠 (백업용)
            if not self.online.active_engines:
                app.update_status_text("Online unavailable. Checking Local AI...")
                self.local.check_health(app)
        else: # Local First
            app.update_status_text("Checking Local AI...")
            self.local.check_health(app)
            # 로컬이 없으면 온라인 체크
            if not self.local.is_available:
                app.update_status_text("Local AI unavailable. Checking Online...")
                self.online.check_health(app)

        app.stop_checking_animation()

    def translate(self, text, task_index, app, logger, task_id):
        # 1. Try Primary
        if self.priority == "online":
            res = self.online.translate(text, task_index, app, logger, task_id)
            if res: return res
            
            # 2. Fallback to Secondary (Local AI)
            if self.local.is_available:
                if app.debug_mode: logger.add(task_id, "FALLBACK", "To_Local", text, "Online Failed")
                return self.local.translate(text, task_index, app, logger, task_id)
                
        else: # Local First
            res = self.local.translate(text, task_index, app, logger, task_id)
            if res: return res
            
            # 2. Fallback to Secondary (Online)
            if app.debug_mode: logger.add(task_id, "FALLBACK", "To_Online", text, "Local Failed")
            return self.online.translate(text, task_index, app, logger, task_id)
            
        return None

    def recover_batch(self, text):
        # 복구 시도: 무조건 둘 다 시도해서 먼저 되는 거 리턴
        res, eng = self.online.recover_batch(text)
        if res: return res, eng
        
        res, eng = self.local.recover_batch(text)
        if res: return res, eng
        
        return None, None

# ===== [Configuration: Active Backend] =====
# 이제 단일 Backend가 아니라 Hybrid Manager를 사용
CURRENT_BACKEND = HybridBackendManager()

# ===== [Logic - Core Processing] =====
# 기존 로직과 100% 동일

def check_engine_health(app):
    CURRENT_BACKEND.check_health(app)

def translate_logic(text, task_index, app, logger, task_id):
    return CURRENT_BACKEND.translate(text, task_index, app, logger, task_id)

def aggressive_recovery_translate(text):
    return CURRENT_BACKEND.recover_batch(text)

def smart_translate(task_info, app, logger):
    task_id = task_info['id']
    lifecycle_manager.update_status(task_id, "IN_PROGRESS")
    try:
        text = task_info['text']
        idx = task_info['index']
        text = text.strip()
        if not text: 
            lifecycle_manager.update_status(task_id, "SKIPPED")
            return None
        if not is_korean_present(text): 
            if app.debug_mode: logger.add(task_id, "SKIPPED", "-", text, "(No Korean)")
            lifecycle_manager.update_status(task_id, "SKIPPED")
            return None
        if is_already_translated_strict(text): 
            if app.debug_mode: logger.add(task_id, "SKIPPED", "-", text, "(Already Translated)")
            lifecycle_manager.update_status(task_id, "SKIPPED")
            return None 

        pattern = r"^\s*([가-하ㄱ-ㅎ①-⑮])(\.|(?:\))|(?:\s))\s+(.*)"
        match = re.match(pattern, text)
        text_to_translate = text
        if match:
            bullet_char = match.group(1)
            content = text[len(bullet_char):].strip()
            if content.startswith(".") or content.startswith(")"): content = content[1:].strip()
            if bullet_char in HAN_TO_ENG_MAP:
                eng_bullet = HAN_TO_ENG_MAP[bullet_char]
                text_to_translate = f"{eng_bullet}. {content}"

        result = translate_logic(text_to_translate, idx, app, logger, task_id)
        if result:
            lifecycle_manager.update_status(task_id, "SUCCESS", result)
            if app.debug_mode: app.log_message(f"[ID:{task_id}] 1st Attempt Success")
        else:
            lifecycle_manager.update_status(task_id, "FAILED")
            if app.debug_mode: app.log_message(f"[ID:{task_id}] 1st Attempt Failed -> Queued", "WARN")
    except Exception as e:
        logger.add(task_id, "ERROR", "CRASH", text, str(e))
        lifecycle_manager.update_status(task_id, "FAILED")

def run_process_thread(input_path, app):
    filename = os.path.basename(input_path)
    try: doc = Document(input_path)
    except Exception as e:
        app.log_message(f"File Open Error ({filename}): {e}", "FATAL")
        return None

    log_file_path = os.path.join(os.path.dirname(input_path), f"log_{filename}.txt")
    logger = FileLogger(log_file_path)
    global lifecycle_manager
    lifecycle_manager = LifecycleManager()
    
    tasks = []
    seen = set()
    counter = 1

    def collect_task(para, is_tbl):
        nonlocal counter
        pid = para._element
        if pid in seen: return
        seen.add(pid)
        if para.text.strip():
            task_id = counter
            lifecycle_manager.register(task_id, para.text)
            tasks.append({'obj': para, 'text': para.text, 'is_table': is_tbl, 'index': len(tasks), 'id': task_id})
            counter += 1

    for para in doc.paragraphs: collect_task(para, False)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs: collect_task(para, True)
    
    total = len(tasks)
    app.log_message(f"[{filename}] Analysis done: {total} items.")
    app.update_progress(0, total, filename)
    
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_to_idx = {executor.submit(smart_translate, t, app, logger): i for i, t in enumerate(tasks)}
        completed = 0
        for future in future_to_idx:
            try: future.result() 
            except: pass
            completed += 1
            app.update_progress(completed, total, filename)

    failed_items = lifecycle_manager.get_failed_tasks()
    if failed_items:
        app.log_message(f"🚨 [{filename}] {len(failed_items)} items failed. Recovery started...", "WARN")
        for tid, orig_text in failed_items:
            res, eng = aggressive_recovery_translate(orig_text)
            if res:
                lifecycle_manager.update_status(tid, "SUCCESS", res)
                logger.add(tid, "RECOVERED", f"{eng}(Recovery)", orig_text, res)
            else:
                logger.add(tid, "FINAL_FAIL", "All", orig_text, "FINAL FAIL")
    
    app.log_message(f"[{filename}] Saving file...")
    saved_log_path = logger.save()
    
    for task in tasks:
        task_id = task['id']
        info = lifecycle_manager.tracking_table.get(task_id)
        if info and info["status"] == "SUCCESS" and info["result"]:
            res = info["result"]
            para = task['obj']
            is_table = task['is_table']
            if is_already_translated_strict(para.text): continue
            if is_table:
                run = para.add_run(f"\n{res}")
                run.italic = True
                run.font.color.rgb = RGBColor(*APPEND_COLOR)
                run.font.size = Pt(8)
            else:
                run = para.add_run(f" ({res})")
                run.italic = True
                run.font.color.rgb = RGBColor(*APPEND_COLOR)
    
    summary = lifecycle_manager.get_summary()
    out_path = get_unique_filename(input_path, "Translated")
    doc.save(out_path)
    app.log_message(f"✅ [{filename}] Done!", "SUCCESS")
    
    app.insert_clickable_path(f"DOC: {os.path.abspath(out_path)}")
    if app.debug_mode or summary['FAILED'] > 0:
        app.insert_clickable_path(f"LOG: {os.path.abspath(saved_log_path)}")
    
    return out_path, log_file_path, summary

# ===== [GUI App] =====
class App:
    def __init__(self, root):
        self.root = root
        self.root.title(f"{APP_NAME} {CURRENT_VERSION}")
        self.root.geometry("850x700")
        self.update_manager = UpdateManager(self)

        saved_debug = config.get("debug_mode", False)
        self.debug_var = tk.BooleanVar(value=saved_debug) 
        self.debug_mode = saved_debug 

        saved_theme = config.get("theme", "light")
        is_dark_init = (saved_theme == "dark")
        self.dark_mode_var = tk.BooleanVar(value=is_dark_init) 

        self.file_paths = []
        self.is_checking_engines = False
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

        # Header
        self.header_frame = tk.Frame(root, pady=15, padx=20, bg="#f8f9fa")
        self.header_frame.pack(fill='x')
        
        # Title & Subtitle Container
        title_container = tk.Frame(self.header_frame, bg="#f8f9fa")
        title_container.pack(side='left')
        
        self.title_lbl = tk.Label(title_container, text=APP_NAME, font=("Segoe UI", 18, "bold"), bg="#f8f9fa", fg="#2c3e50")
        self.title_lbl.pack(side='left')
        
        # Subtitle Label
        self.subtitle_lbl = tk.Label(title_container, text=APP_SUBTITLE, font=("Segoe UI", 10), bg="#f8f9fa", fg="#7f8c8d")
        self.subtitle_lbl.pack(side='left', padx=(10, 0), pady=(8, 0))

        self.update_btn = tk.Button(self.header_frame, text="Up to date", state='disabled', relief='flat', bg="#e9ecef")
        self.update_btn.pack(side='right')

        # Main Body
        self.main_frame = ttk.Frame(root, padding="20")
        self.main_frame.pack(fill='both', expand=True)
        
        # Controls Row
        ctl_frame = ttk.Frame(self.main_frame)
        ctl_frame.pack(fill='x', pady=5)
        self.chk_debug = ttk.Checkbutton(ctl_frame, text="Debug Mode", variable=self.debug_var, command=self.toggle_debug)
        self.chk_debug.pack(side='left', padx=5)
        self.chk_dark = ttk.Checkbutton(ctl_frame, text="Dark Mode", variable=self.dark_mode_var, command=self.toggle_theme)
        self.chk_dark.pack(side='left', padx=5)

        # Buttons
        self.btn_select = ttk.Button(self.main_frame, text="📂 Select .docx Files", command=self.select_files)
        self.btn_select.pack(pady=15, fill='x', ipady=8)
        self.btn_run = ttk.Button(self.main_frame, text="🚀 Start Translation", command=self.start_thread, state='disabled')
        self.btn_run.pack(pady=5, fill='x', ipady=8)
        
        # Status
        self.lbl_status_detail = ttk.Label(self.main_frame, text="Please select Word files to start.", font=("Segoe UI", 10))
        self.lbl_status_detail.pack(pady=(15, 5))
        self.progress = ttk.Progressbar(self.main_frame, length=200, mode="determinate")
        self.progress.pack(pady=5, fill='x')
        
        # Log Area
        self.log_frame = ttk.LabelFrame(self.main_frame, text="Process Log", padding="10")
        self.log_frame.pack(fill="both", expand=True, pady=10)
        
        self.log_area = scrolledtext.ScrolledText(self.log_frame, state='disabled', cursor="arrow", font=("Consolas", 9))
        self.log_area.pack(fill="both", expand=True)
        
        self.log_area.tag_config("WARN", foreground="orange")
        self.log_area.tag_config("SUCCESS", foreground="green")
        self.log_area.tag_config("FATAL", foreground="red", font=("Segoe UI", 10, "bold"))
        self.log_area.tag_config("HYPERLINK", foreground="blue", underline=True)
        self.log_area.tag_bind("HYPERLINK", "<Button-1>", self.on_link_click)
        self.log_area.tag_bind("HYPERLINK", "<Enter>", lambda e: self.log_area.config(cursor="hand2"))
        self.log_area.tag_bind("HYPERLINK", "<Leave>", lambda e: self.log_area.config(cursor="arrow"))
        
        if is_dark_init:
            self.toggle_theme()
            
        if saved_debug:
            self.toggle_debug()

        threading.Thread(target=self.update_manager.check_for_updates, daemon=True).start()

    def insert_clickable_path(self, text):
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, text + "\n", "HYPERLINK")
        self.log_area.insert(tk.END, "(Click to open)\n", "gray")
        self.log_area.tag_config("gray", foreground="gray", font=("Segoe UI", 8))
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')

    def on_link_click(self, event):
        try:
            index = self.log_area.index(f"@{event.x},{event.y}")
            line_text = self.log_area.get(f"{index.split('.')[0]}.0", f"{index.split('.')[0]}.end").strip()
            if "DOC: " in line_text: open_file_safe(line_text.replace("DOC: ", ""))
            elif "LOG: " in line_text: open_file_safe(line_text.replace("LOG: ", ""))
        except: pass

    def start_checking_animation(self):
        self.is_checking_engines = True
        threading.Thread(target=self._animate_checking, daemon=True).start()

    def _animate_checking(self):
        dots = ["", ".", "..", "..."]
        idx = 0
        while self.is_checking_engines:
            msg = f"Checking translation engines{dots[idx % 4]}"
            self.update_status_text(msg)
            idx += 1
            time.sleep(0.5)
    
    def stop_checking_animation(self):
        self.is_checking_engines = False
        time.sleep(0.6)
        self.update_status_text("Done! Engines ready.")
        time.sleep(1)

    def update_status_text(self, text):
        self.root.after(0, lambda: self.lbl_status_detail.config(text=text))

    def toggle_theme(self):
        is_dark = self.dark_mode_var.get()
        # Header Colors
        bg_header = "#3c3f41" if is_dark else "#f8f9fa"
        fg_title = "white" if is_dark else "#2c3e50"
        fg_sub = "#bbbbbb" if is_dark else "#7f8c8d"
        
        # Log Colors
        bg_log = "#1e1e1e" if is_dark else "white"
        fg_log = "#d4d4d4" if is_dark else "black"
        
        # Apply
        self.header_frame.config(bg=bg_header)
        self.title_lbl.config(bg=bg_header, fg=fg_title)
        self.subtitle_lbl.config(bg=bg_header, fg=fg_sub) 
        self.title_lbl.master.config(bg=bg_header) 
        
        self.log_area.config(bg=bg_log, fg=fg_log, insertbackground="white" if is_dark else "black")
        
        if is_dark:
            self.log_area.tag_config("HYPERLINK", foreground="#61afef")
            config.set("theme", "dark")
        else:
            self.log_area.tag_config("HYPERLINK", foreground="blue")
            config.set("theme", "white")

    def notify_update(self, new_ver, msg):
        def flash_alert():
            colors = ["#ffcccc", "#ff9999", "#ff6666", "#ff3333"]
            idx = 0
            while self.update_manager.update_available:
                try:
                    self.update_btn.config(bg=colors[idx % 4])
                    idx += 1
                    time.sleep(0.5)
                except: break
        self.update_btn.config(text=f"🚨 Update ({new_ver})", state='normal', 
                               command=lambda: self.show_update_dialog(new_ver, msg), fg="black")
        ToolTip(self.update_btn, f"New version {new_ver} available!\n{msg}")
        threading.Thread(target=flash_alert, daemon=True).start()

    def show_update_dialog(self, new_ver, msg):
        dialog = tk.Toplevel(self.root)
        dialog.title("Update")
        dialog.geometry("400x300")
        tk.Label(dialog, text=f"New Version {new_ver}!", font=("Segoe UI", 12, "bold"), fg="blue").pack(pady=15)
        msg_area = tk.Text(dialog, height=5, width=40); msg_area.insert("1.0", msg); msg_area.config(state='disabled'); msg_area.pack(pady=5)
        def go_dl(): webbrowser.open(RELEASE_URL); dialog.destroy()
        def skip(): 
            config.set("ignored_version", new_ver)
            self.update_manager.update_available = False
            self.update_btn.config(text="Skipped", state='disabled', bg="#e9ecef")
            dialog.destroy()
        tk.Button(dialog, text="Download", command=go_dl, bg="#ddddff").pack(pady=5)
        tk.Button(dialog, text="Skip", command=skip).pack(pady=5)

    def toggle_debug(self):
        self.debug_mode = self.debug_var.get()
        config.set("debug_mode", self.debug_mode)
        if self.debug_mode: self.log_message("🕵️ Debug Mode ON")
        else: self.log_message("🚀 High-Speed Mode")

    def log_message(self, msg, tag=None):
        if not self.debug_mode and tag not in ["SUCCESS", "WARN", "FATAL"]: return
        def _log():
            self.log_area.config(state='normal')
            self.log_area.insert(tk.END, f"{msg}\n", tag)
            self.log_area.see(tk.END)
            self.log_area.config(state='disabled')
        self.root.after(0, _log)

    def select_files(self):
        files = filedialog.askopenfilenames(filetypes=[("Word files", "*.docx")])
        if files:
            self.file_paths = files
            self.log_message(f"{len(files)} files selected.")
            self.update_status_text(f"Ready: {len(files)} files selected.")
            self.btn_run.config(state='normal')

    def start_thread(self):
        if not self.file_paths: return
        self.btn_select.config(state='disabled')
        self.btn_run.config(state='disabled')
        self.debug_mode = self.debug_var.get()
        self.log_area.config(state='normal')
        self.log_area.delete(1.0, tk.END) 
        self.log_area.config(state='disabled')
        t = threading.Thread(target=self.run_batch_logic)
        t.start()
        
    def run_batch_logic(self):
        check_engine_health(self)
        total_files = len(self.file_paths)
        success_files = 0
        for i, path in enumerate(self.file_paths):
            current_num = i + 1
            filename = os.path.basename(path)
            self.current_file_info = f"[{current_num}/{total_files}] {filename}"
            self.update_progress(0, 100, filename)
            self.log_message(f"=== Processing {self.current_file_info} ===", "SUCCESS")
            res = run_process_thread(path, self)
            if res: success_files += 1
            time.sleep(1)
        messagebox.showinfo("Done", f"All tasks finished!\nSuccess: {success_files}/{total_files}")
        self.reset_ui()
            
    def update_progress(self, curr, total, filename=""):
        file_idx_info = getattr(self, 'current_file_info', "")
        if total > 0:
            pct = (curr / total) * 100
            msg = f"{file_idx_info} - {int(pct)}% ({curr}/{total})"
            self.root.after(0, lambda: self.progress.configure(value=pct))
            self.root.after(0, lambda: self.lbl_status_detail.config(text=msg))

    def reset_ui(self):
        self.file_paths = [] 
        self.root.after(0, lambda: self.lbl_status_detail.config(text="Please select Word files (.docx)"))
        self.root.after(0, lambda: self.progress.configure(value=0))
        self.root.after(0, lambda: self.btn_select.config(state='normal'))
        self.root.after(0, lambda: self.btn_run.config(state='disabled')) 

    def on_closing(self):
        try:
            config.save() 
        except:
            pass
        self.root.destroy() 

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()